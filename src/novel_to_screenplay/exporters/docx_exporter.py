"""Render a screenplay document as a .docx (Word) file.

Chinese screenwriting workflows are overwhelmingly Word-based, so a clean
.docx with a conventional layout (scene headings, 人物：对白, transitions) is
often the most directly useful deliverable for the local market.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

LOCATION_PREFIXES = {
    "INT": "INT.",
    "EXT": "EXT.",
    "INT_EXT": "INT./EXT.",
}
TIME_OF_DAY = {
    "DAY": "日",
    "NIGHT": "夜",
    "MORNING": "晨",
    "EVENING": "黄昏",
    "CONTINUOUS": "连续",
    "LATER": "稍后",
}


def write_docx(document: dict[str, Any], output_path: Path) -> None:
    """Write the screenplay document to a .docx file at output_path."""

    metadata = document.get("metadata")
    metadata = metadata if isinstance(metadata, dict) else {}

    doc = Document()
    title = str(metadata.get("title", "剧本初稿")).strip() or "剧本初稿"
    doc.add_heading(title, level=0)
    author = metadata.get("author")
    if isinstance(author, str) and author.strip():
        doc.add_paragraph(f"编剧：{author.strip()}")

    for scene in _scenes(document):
        heading = doc.add_paragraph()
        heading.add_run(_scene_heading_text(scene)).bold = True
        for element in _script(scene):
            _add_element(doc, element)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def _scene_heading_text(scene: dict[str, Any]) -> str:
    heading = scene.get("scene_heading", {})
    heading = heading if isinstance(heading, dict) else {}
    display = str(heading.get("display", "未知地点")).strip() or "未知地点"
    prefix = LOCATION_PREFIXES.get(heading.get("location_mode", ""), "")
    time = TIME_OF_DAY.get(heading.get("time_of_day", ""), "")
    location = f"{prefix} {display}".strip()
    return f"{location} - {time}" if time else location


def _add_element(doc: Document, element: dict[str, Any]) -> None:
    element_type = element.get("type")
    text = str(element.get("text", "")).strip()
    if not text:
        return

    if element_type in {"dialogue", "voiceover"}:
        name = str(element.get("character_name", "")).strip()
        suffix = "（V.O.）" if element_type == "voiceover" else ""
        paragraph = doc.add_paragraph()
        if name:
            paragraph.add_run(f"{name}{suffix}：").bold = True
        parenthetical = element.get("parenthetical")
        if isinstance(parenthetical, str) and parenthetical.strip():
            paragraph.add_run(f"（{parenthetical.strip()}）")
        paragraph.add_run(text)
        subtext = element.get("subtext")
        if isinstance(subtext, str) and subtext.strip():
            doc.add_paragraph().add_run(f"潜台词：{subtext.strip()}").italic = True
        return

    if element_type == "transition":
        paragraph = doc.add_paragraph(text)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        return

    if element_type in {"note", "parenthetical"}:
        doc.add_paragraph().add_run(f"【批注】{text}").italic = True
        return

    doc.add_paragraph(text)


def _scenes(document: dict[str, Any]) -> list[dict[str, Any]]:
    scenes = document.get("scenes")
    if not isinstance(scenes, list):
        return []
    return [scene for scene in scenes if isinstance(scene, dict)]


def _script(scene: dict[str, Any]) -> list[dict[str, Any]]:
    script = scene.get("script")
    if not isinstance(script, list):
        return []
    return [element for element in script if isinstance(element, dict)]
